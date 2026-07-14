# -*- coding: utf-8 -*-
"""
DOCX 合同文本提取模块（python-docx 引擎）

技术方案：
  - 使用 python-docx 读取段落（paragraphs）+ 表格（tables）
  - 段落用换行拼接，表格每行单元格以制表符分隔、行之间换行
  - 不再依赖 pdfplumber 兜底，DOCX 走独立解析链路

输入：DOCX / DOC 文件路径（.doc 需提前转为 docx，本函数只处理 .docx）
输出：提取后的全文文本字符串（保留段落结构）
"""

import os
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def parse_docx(file_path: str) -> str:
    """
    提取 DOCX 合同全文，段落与表格按阅读顺序拼接。

    Args:
        file_path: DOCX 文件路径

    Returns:
        str: 合同全文文本

    Raises:
        FileNotFoundError: 文件不存在
        ValueError:      解析失败（非 docx / 损坏）
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX 文件不存在: {file_path}")

    try:
        doc = Document(file_path)
    except PackageNotFoundError as e:
        raise ValueError(f"文件不是有效的 docx 格式: {file_path}") from e
    except Exception as e:
        raise ValueError(f"无法打开 docx 文件 {file_path}: {e}") from e

    parts = []

    # 1. 段落文本
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    # 2. 表格文本：逐行逐列，制表符分隔
    for table in doc.tables:
        try:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                line = "\t".join(c for c in cells if c)
                if line:
                    parts.append(line)
        except Exception:
            # 兼容合并单元格等异常结构，单表失败不中断
            continue

    if not parts:
        raise ValueError(f"DOCX 未提取到任何文本: {file_path}")

    return "\n".join(parts)
